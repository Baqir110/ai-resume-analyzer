import io
from typing import List, Optional
from pydantic import BaseModel

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from fastapi import (
    APIRouter,
    File,
    Form,
    HTTPException,
    UploadFile,
)
from fastapi.responses import StreamingResponse

from app.models.schemas import AnalysisResponse
from app.services.analyzer import analyze_resume_content
from app.services.audit_matrix import AuditMatrixService
from app.services.bulk_analyzer import BulkAnalyzerService
from app.services.cover_letter import CoverLetterService
from app.services.diff_preview import DiffPreviewService
from app.services.interview_prep import InterviewPrepService
from app.services.latex_generator import (
    compile_latex_to_pdf,
    generate_german_latex_content,
)
from app.services.linkedin_optimizer import LinkedInOptimizerService
from app.services.llm_provider import (
    LOG_PATH,
    LLMService,
)
from app.services.optimizer import (
    generate_full_tailored_cv,
    optimize_resume_bullets,
    suggest_best_cv_format,
)
from app.services.parser import extract_text_from_file
from app.services.tracker import ApplicationTrackerService

router = APIRouter()


# ============================================================
# PUSHED FLAGSHIP SCHEMAS & MODELS
# ============================================================


class BulletDiffRequest(BaseModel):
    original_bullets: List[str]
    optimized_bullets: List[str]


class TrackerCreateRequest(BaseModel):
    company_name: str
    job_title: str
    job_url: Optional[str] = ""
    ats_score: Optional[int] = 0
    status: Optional[str] = "Saved"
    notes: Optional[str] = ""


class TrackerStatusUpdate(BaseModel):
    status: str


# ============================================================
# FLAGSHIP 1: VISUAL DIFF & BULK ANALYSIS
# ============================================================


@router.post("/diff-preview")
async def diff_preview(payload: BulletDiffRequest):
    try:
        diffs = DiffPreviewService.compare_bullet_lists(
            payload.original_bullets, payload.optimized_bullets
        )
        return {"status": "success", "diffs": diffs}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@router.post("/analyze-bulk")
async def analyze_bulk(
    job_description: str = Form(...),
    resume_files: List[UploadFile] = File(...),
):
    if not resume_files:
        raise HTTPException(status_code=400, detail="No resume files uploaded.")

    try:
        files_data = []
        for file in resume_files:
            content = await file.read()
            files_data.append((content, file.filename))

        ranked_candidates = await BulkAnalyzerService.process_batch(
            files_data, job_description
        )

        return {
            "status": "success",
            "total_processed": len(ranked_candidates),
            "rankings": ranked_candidates,
        }
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


# ============================================================
# FLAGSHIP 2: ATS AUDIT MATRIX
# ============================================================


@router.post("/audit-matrix")
async def audit_matrix_endpoint(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
):
    resume_text = await extract_text_from_file(resume_file)
    if not resume_text:
        raise HTTPException(status_code=400, detail="Could not extract resume text.")

    file_ext = resume_file.filename.split(".")[-1] if resume_file.filename else "pdf"
    audit_result = AuditMatrixService.run_full_audit(
        resume_text=resume_text,
        job_description=job_description,
        file_type=file_ext,
    )
    return {"status": "success", "data": audit_result}


# ============================================================
# FLAGSHIP 3: COVER LETTER & OUTREACH
# ============================================================


@router.post("/generate-cover-letter")
async def generate_cover_letter_endpoint(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    company_name: Optional[str] = Form("Target Company"),
    tone: Optional[str] = Form("formal"),
    provider: Optional[str] = Form("gemini"),
):
    resume_text = await extract_text_from_file(resume_file)
    if not resume_text:
        raise HTTPException(status_code=400, detail="Could not extract resume text.")

    result = CoverLetterService.generate_cover_letter_and_outreach(
        resume_text=resume_text,
        job_description=job_description,
        company_name=company_name or "Target Company",
        tone=tone or "formal",
        provider=provider or "gemini",
    )
    return {"status": "success", "data": result}


# ============================================================
# FLAGSHIP 4: INTERVIEW PREP & GAP DEFENSE
# ============================================================


@router.post("/interview-prep")
async def interview_prep_endpoint(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    provider: Optional[str] = Form("gemini"),
):
    resume_text = await extract_text_from_file(resume_file)
    if not resume_text:
        raise HTTPException(status_code=400, detail="Could not extract resume text.")

    analysis = analyze_resume_content(resume_text, job_description)
    missing_skills = analysis.get("missing_skills") or []

    prep_data = InterviewPrepService.generate_interview_prep(
        resume_text=resume_text,
        job_description=job_description,
        missing_skills=missing_skills,
        provider=provider or "gemini",
    )
    return {"status": "success", "data": prep_data}


# ============================================================
# FLAGSHIP 5: LINKEDIN PROFILE OPTIMIZER
# ============================================================


@router.post("/linkedin-optimize")
async def linkedin_optimize_endpoint(
    resume_file: UploadFile = File(...),
    target_role: Optional[str] = Form("Software Engineer"),
    provider: Optional[str] = Form("gemini"),
):
    resume_text = await extract_text_from_file(resume_file)
    if not resume_text:
        raise HTTPException(status_code=400, detail="Could not extract resume text.")

    optimized = LinkedInOptimizerService.optimize_profile(
        resume_text=resume_text,
        target_role=target_role or "Software Engineer",
        provider=provider or "gemini",
    )
    return {"status": "success", "data": optimized}


# ============================================================
# FLAGSHIP 6: APPLICATION PIPELINE TRACKER
# ============================================================


@router.get("/tracker/applications")
async def list_applications_endpoint():
    return {
        "status": "success",
        "applications": ApplicationTrackerService.list_applications(),
    }


@router.post("/tracker/applications")
async def create_application_endpoint(payload: TrackerCreateRequest):
    app_data = ApplicationTrackerService.create_application(
        company_name=payload.company_name,
        job_title=payload.job_title,
        job_url=payload.job_url,
        ats_score=payload.ats_score or 0,
        status=payload.status or "Saved",
        notes=payload.notes or "",
    )
    return {"status": "success", "data": app_data}


@router.patch("/tracker/applications/{app_id}")
async def update_status_endpoint(app_id: int, payload: TrackerStatusUpdate):
    updated = ApplicationTrackerService.update_status(app_id, payload.status)
    if not updated:
        raise HTTPException(status_code=404, detail="Application record not found.")
    return {"status": "success", "message": "Application status updated."}


@router.delete("/tracker/applications/{app_id}")
async def delete_application_endpoint(app_id: int):
    deleted = ApplicationTrackerService.delete_application(app_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Application record not found.")
    return {"status": "success", "message": "Application deleted."}


# ============================================================
# DOCX BUILDER
# ============================================================


def build_ats_docx_resume(markdown_resume: str) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    for line in markdown_resume.splitlines():
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(18)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif stripped.startswith("## "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(stripped[3:].upper())
            run.bold = True
            run.font.size = Pt(12)

        elif stripped.startswith(("* ", "- ")):
            paragraph = doc.add_paragraph(
                stripped[2:].strip(),
                style="List Bullet",
            )
            paragraph.style.font.size = Pt(10)

        else:
            paragraph = doc.add_paragraph(stripped)
            paragraph.style.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# LIGHTWEIGHT KEEP-ALIVE HEALTH ENDPOINT
# ============================================================


@router.get("/health")
async def health_check():
    """Lightweight endpoint for keep-alive pings."""
    return {"status": "ok"}


# ============================================================
# ANALYZE RESUME
# ============================================================


@router.post(
    "/analyze",
    response_model=AnalysisResponse,
)
async def analyze_resume(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    provider: Optional[str] = Form("gemini"),
    model_name: Optional[str] = Form(None),
):
    if not job_description.strip():
        raise HTTPException(
            status_code=400,
            detail="Job description cannot be empty.",
        )

    resume_text = await extract_text_from_file(resume_file)

    if not resume_text:
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from resume.",
        )

    results = analyze_resume_content(
        resume_text=resume_text,
        job_description=job_description,
    )

    results["recommendation"] = suggest_best_cv_format(
        job_description=job_description,
        resume_text=resume_text,
    )

    if results.get("missing_skills"):
        try:
            print(
                f"[AI DEBUG] /analyze "
                f"provider={provider!r} "
                f"model_name={model_name!r} "
                f"missing_skills={results.get('missing_skills')!r}"
            )

            rewrite = optimize_resume_bullets(
                resume_text=resume_text,
                job_description=job_description,
                missing_skills=results["missing_skills"],
                provider=provider or "gemini",
            )

            results.setdefault(
                "improvement_suggestions",
                [],
            ).append(
                f"AI Bullet Point Rewrite "
                f"({(provider or 'gemini').upper()}):\n\n"
                f"{rewrite}"
            )

        except Exception as exc:
            print(f"[AI DEBUG] /analyze bullet rewrite failed: {exc}")

            results.setdefault(
                "improvement_suggestions",
                [],
            ).append(f"AI bullet rewrite unavailable: {exc}")

    return AnalysisResponse(
        status="success",
        ats_match_score=results["ats_match_score"],
        keyword_density_score=results["keyword_density_score"],
        matching_skills=results["matching_skills"],
        missing_skills=results["missing_skills"],
        improvement_suggestions=results["improvement_suggestions"],
        recommendation=results.get("recommendation"),
    )


# ============================================================
# FULL TAILORED DOCX GENERATION
# ============================================================


@router.post("/generate-full")
async def generate_full_cv_endpoint(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    provider: Optional[str] = Form("gemini"),
    model_name: Optional[str] = Form(None),
):
    resume_text = await extract_text_from_file(resume_file)

    if not resume_text:
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from resume.",
        )

    analysis = analyze_resume_content(
        resume_text=resume_text,
        job_description=job_description,
    )

    missing_skills = analysis.get("missing_skills") or []

    tailored = generate_full_tailored_cv(
        resume_text=resume_text,
        job_description=job_description,
        missing_skills=missing_skills,
        provider=provider or "gemini",
    )

    docx_bytes = build_ats_docx_resume(tailored)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": (
                "attachment; filename=Tailored_Optimized_Resume.docx"
            )
        },
    )


# ============================================================
# GERMAN PDF LEBENSLAUF
# ============================================================


@router.post("/generate-german-cv")
async def generate_german_cv_endpoint(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    layout_style: Optional[str] = Form("german_corporate"),
    template_style: Optional[str] = Form(None),
    provider: Optional[str] = Form("gemini"),
    model_name: Optional[str] = Form(None),
):
    selected_style = template_style or layout_style or "german_corporate"

    resume_text = await extract_text_from_file(resume_file)

    if not resume_text:
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from resume.",
        )

    analysis = analyze_resume_content(
        resume_text=resume_text,
        job_description=job_description,
    )

    missing_skills = analysis.get("missing_skills") or []

    latex_code = generate_german_latex_content(
        resume_text=resume_text,
        job_description=job_description,
        missing_skills=missing_skills,
        provider=provider or "gemini",
        model_name=model_name,
        layout_style=selected_style,
    )

    try:
        pdf_bytes = compile_latex_to_pdf(latex_code)
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"LaTeX compilation failed.\n\n{exc}",
        ) from exc

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": (
                "attachment; "
                f"filename=Lebenslauf_Muhammad_Baqir_{selected_style}.pdf"
            )
        },
    )


# ============================================================
# GERMAN LATEX TEX SOURCE
# ============================================================


@router.post("/generate-tex-cv")
async def generate_tex_cv_endpoint(
    job_description: str = Form(...),
    resume_file: UploadFile = File(...),
    layout_style: Optional[str] = Form("german_corporate"),
    template_style: Optional[str] = Form(None),
    provider: Optional[str] = Form("gemini"),
    model_name: Optional[str] = Form(None),
):
    selected_style = template_style or layout_style or "german_corporate"

    resume_text = await extract_text_from_file(resume_file)

    if not resume_text:
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from resume.",
        )

    analysis = analyze_resume_content(
        resume_text=resume_text,
        job_description=job_description,
    )

    missing_skills = analysis.get("missing_skills") or []

    latex_code = generate_german_latex_content(
        resume_text=resume_text,
        job_description=job_description,
        missing_skills=missing_skills,
        provider=provider or "gemini",
        model_name=model_name,
        layout_style=selected_style,
    )

    return StreamingResponse(
        io.BytesIO(latex_code.encode("utf-8")),
        media_type="text/plain",
        headers={
            "Content-Disposition": (
                "attachment; "
                f"filename=Lebenslauf_Muhammad_Baqir_{selected_style}.tex"
            )
        },
    )


# ============================================================
# BACKEND LLM STATUS & LOGGING
# ============================================================


@router.get("/backend-status")
async def backend_status():
    return {
        "status": "online",
        "providers": LLMService.provider_status(),
        "log_file": str(LOG_PATH),
    }


@router.get("/processing-log")
async def processing_log(limit: int = 100):
    limit = max(1, min(limit, 500))
    logs = LLMService.recent_logs(limit)

    return {
        "status": "success",
        "count": len(logs),
        "logs": logs,
    }


@router.delete("/processing-log")
async def clear_processing_log():
    LLMService.clear_logs()

    return {
        "status": "success",
        "message": "Backend processing log cleared.",
    }