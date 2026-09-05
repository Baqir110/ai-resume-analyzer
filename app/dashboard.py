import io
import os
import time
from pathlib import Path

import requests
import streamlit as st
from docx import Document

try:
    from app.services.llm_provider import LLMService, LOG_PATH
except Exception:
    LLMService = None
    LOG_PATH = Path(
        os.getenv(
            "LLM_PROCESSING_LOG",
            str(Path.home() / ".resume_hub_llm.log.jsonl"),
        )
    )


# ==============================================================================
# CONFIGURATION
# ==============================================================================

# Dynamic base URL: Checks environment variable first, then Render, then local fallback
DEFAULT_API_BASE = os.getenv(
    "FASTAPI_API_BASE",
    "https://ai-resume-backend-vowl.onrender.com",
)
REQUEST_TIMEOUT = 120

PROVIDER_LABELS = {
    "gemini": "Google Gemini",
    "groq": "Groq",
    "openrouter": "OpenRouter",
    "deepseek": "DeepSeek",
    "ollama": "Ollama",
    "openai": "OpenAI",
    "claude": "Anthropic Claude",
}

DEFAULT_MODELS = {
    "gemini": "gemini-2.5-flash",
    "groq": "openai/gpt-oss-120b",
    "openrouter": "deepseek/deepseek-chat",
    "deepseek": "deepseek-chat",
    "ollama": "llama3",
    "openai": "gpt-4o-mini",
    "claude": "claude-3-5-haiku-20241022",
}

TEMPLATE_LABELS = {
    "german_corporate": "Corporate Slate Navy",
    "german_modern": "Modern Two-Column",
    "german_classic": "Classic Conservative",
    "international_ats": "International English ATS",
}


# ==============================================================================
# PAGE CONFIG
# ==============================================================================

st.set_page_config(
    page_title="AI Resume & CV Optimization Hub",
    page_icon="🎯",
    layout="wide",
)


# ==============================================================================
# API ENDPOINTS
# ==============================================================================


def endpoints(api_base: str) -> dict:
    """FastAPI routes definition."""
    api_base = api_base.rstrip("/")
    prefix = f"{api_base}/api/v1/resume"

    return {
        "analyze": f"{prefix}/analyze",
        "full_docx": f"{prefix}/generate-full",
        "german_pdf": f"{prefix}/generate-german-cv",
        "german_tex": f"{prefix}/generate-tex-cv",
    }


# ==============================================================================
# SESSION STATE
# ==============================================================================

if "api_base" not in st.session_state:
    st.session_state["api_base"] = DEFAULT_API_BASE

if "provider" not in st.session_state:
    st.session_state["provider"] = "gemini"

if "last_analysis" not in st.session_state:
    st.session_state["last_analysis"] = None

if "uploaded_file_data" not in st.session_state:
    st.session_state["uploaded_file_data"] = None

if "job_desc" not in st.session_state:
    st.session_state["job_desc"] = ""


# ==============================================================================
# STYLING
# ==============================================================================

st.markdown(
    """
<style>

.chip {
    display: inline-block;
    padding: 5px 12px;
    margin: 3px;
    border-radius: 999px;
    font-size: 0.85rem;
    font-weight: 500;
}

.chip-good {
    background: #DCFCE7;
    color: #166534;
}

.chip-bad {
    background: #FEE2E2;
    color: #991B1B;
}

.step-badge {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 28px;
    height: 28px;
    border-radius: 50%;
    background: #E5E7EB;
    color: #374151;
    font-weight: 600;
    font-size: 0.85rem;
    margin-right: 8px;
}

.step-badge.active {
    background: #2563EB;
    color: white;
}

.step-badge.done {
    background: #16A34A;
    color: white;
}

.metric-big {
    font-size: 2.4rem;
    font-weight: 700;
    line-height: 1;
}

.metric-label {
    color: #6B7280;
    font-size: 0.85rem;
    text-transform: uppercase;
    letter-spacing: 0.04em;
}

.provider-card {
    padding: 10px 12px;
    border-radius: 8px;
    margin-bottom: 6px;
    border: 1px solid #E5E7EB;
}

.security-note {
    padding: 12px;
    border-radius: 8px;
    background: #1E293B;
    border: 1px solid #334155;
    color: #F8FAFC;
}

</style>
""",
    unsafe_allow_html=True,
)


# ==============================================================================
# HELPERS
# ==============================================================================


def chips(items, kind="good"):
    if not items:
        st.caption("None found.")
        return

    css_class = "chip-good" if kind == "good" else "chip-bad"
    html = "".join(f'<span class="chip {css_class}">{item}</span>' for item in items)
    st.markdown(html, unsafe_allow_html=True)


def score_block(label: str, value: int):
    value = int(value or 0)
    value = max(0, min(value, 100))

    st.markdown(
        f'<div class="metric-label">{label}</div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        f'<div class="metric-big">{value}%</div>',
        unsafe_allow_html=True,
    )

    st.progress(value / 100)


def error_detail(response) -> str:
    if response is None:
        return "Could not connect to the backend service."

    try:
        data = response.json()
        if isinstance(data, dict):
            detail = data.get("detail")
            if detail:
                return str(detail)
        return response.text
    except Exception:
        return response.text or f"HTTP {response.status_code}"


def api_request(
    url: str,
    data: dict,
    files: dict,
    spinner_text: str,
):
    with st.status(
        spinner_text,
        expanded=True,
    ) as status:
        st.write("📄 Sending payload and document to backend...")

        try:
            response = requests.post(
                url,
                data=data,
                files=files,
                timeout=REQUEST_TIMEOUT,
            )
        except requests.exceptions.ConnectionError:
            status.update(
                label="Backend unavailable.",
                state="error",
            )
            return None
        except requests.exceptions.Timeout:
            status.update(
                label="Backend request timed out.",
                state="error",
            )
            return None
        except requests.exceptions.RequestException as exc:
            status.update(
                label="Request failed.",
                state="error",
            )
            st.error(str(exc))
            return None

        if response.status_code in (429, 503):
            wait_time = 10 if response.status_code == 429 else 5
            status.update(
                label=(
                    f"Server busy ({response.status_code}). "
                    f"Retrying in {wait_time}s..."
                ),
                state="running",
            )

            progress = st.progress(0)
            for i in range(wait_time):
                time.sleep(1)
                progress.progress((i + 1) / wait_time)
            progress.empty()

            try:
                response = requests.post(
                    url,
                    data=data,
                    files=files,
                    timeout=REQUEST_TIMEOUT,
                )
            except requests.exceptions.RequestException:
                status.update(
                    label="Retry failed.",
                    state="error",
                )
                return None

        if response.status_code == 200:
            status.update(
                label="✅ Analysis complete!",
                state="complete",
                expanded=False,
            )
        else:
            status.update(
                label=f"Request failed ({response.status_code})",
                state="error",
            )

    return response


def build_files_payload():
    uploaded = st.session_state.get("uploaded_file_data")
    if not uploaded:
        return {}

    filename, file_bytes, mime_type = uploaded

    return {
        "resume_file": (
            filename,
            file_bytes,
            mime_type,
        )
    }


def generate_txt_export(result_data: dict) -> bytes:
    lines = [
        "=" * 60,
        "AI RESUME & CV OPTIMIZATION REPORT".center(60),
        "=" * 60,
        "",
        f"ATS Match Score: {result_data.get('ats_match_score', 0)}%",
        f"Keyword Density Score: {result_data.get('keyword_density_score', 0)}%",
        "",
        "[MATCHING SKILLS]",
        ", ".join(result_data.get("matching_skills", [])) or "None",
        "",
        "[MISSING SKILLS]",
        ", ".join(result_data.get("missing_skills", [])) or "None",
        "",
        "=" * 60,
        "SUGGESTIONS & REWRITES".center(60),
        "=" * 60,
        "",
    ]

    for index, suggestion in enumerate(
        result_data.get("improvement_suggestions", []),
        1,
    ):
        lines.append(f"{index}. {suggestion}")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


def generate_docx_export(result_data: dict) -> bytes:
    doc = Document()

    doc.add_heading(
        "AI Resume Optimization Report",
        level=0,
    )

    doc.add_heading(
        "Overview Metrics",
        level=1,
    )

    paragraph = doc.add_paragraph()
    paragraph.add_run("ATS Match Score: ").bold = True
    paragraph.add_run(f"{result_data.get('ats_match_score', 0)}%\n")
    paragraph.add_run("Keyword Density: ").bold = True
    paragraph.add_run(f"{result_data.get('keyword_density_score', 0)}%")

    doc.add_heading(
        "Skill Gap Analysis",
        level=1,
    )

    paragraph = doc.add_paragraph()
    paragraph.add_run("Matching Skills: ").bold = True
    paragraph.add_run(", ".join(result_data.get("matching_skills", [])) or "None")
    paragraph.add_run("\n")
    paragraph.add_run("Missing Skills: ").bold = True
    paragraph.add_run(", ".join(result_data.get("missing_skills", [])) or "None")

    doc.add_heading(
        "Suggestions & Rewritten Bullets",
        level=1,
    )

    for suggestion in result_data.get("improvement_suggestions", []):
        if suggestion.startswith("🤖"):
            doc.add_heading("AI Rewritten Content", level=2)
            doc.add_paragraph(suggestion.replace("🤖 ", ""))
        else:
            doc.add_paragraph(suggestion, style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def reset_all():
    st.session_state["last_analysis"] = None
    st.session_state["uploaded_file_data"] = None
    st.session_state["job_desc"] = ""


# ==============================================================================
# SIDEBAR
# ==============================================================================

with st.sidebar:
    st.title("AI Engine")

    provider_options = list(PROVIDER_LABELS.keys())
    current_provider = st.session_state.get("provider", "gemini")

    provider = st.selectbox(
        "AI Provider",
        options=provider_options,
        index=(
            provider_options.index(current_provider)
            if current_provider in provider_options
            else 0
        ),
        format_func=lambda value: PROVIDER_LABELS[value],
    )

    st.session_state["provider"] = provider
    model_name = DEFAULT_MODELS.get(provider, "")

    st.markdown(
        """
        <div class="security-note">
        <strong>API keys are backend-only.</strong><br><br>
        The dashboard does not request, store, or transmit API keys.
        All provider credentials are loaded from the backend
        <code>.env</code> file.
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.write("")
    st.caption(f"Model: `{model_name}`")

    with st.expander("Backend configuration", expanded=False):
        api_base_input = st.text_input(
            "FastAPI Backend URL",
            value=st.session_state.get("api_base", DEFAULT_API_BASE),
            help="Example: http://localhost:8000 or your Render URL",
        )

        if api_base_input.strip():
            st.session_state["api_base"] = api_base_input.strip().rstrip("/")

        st.caption("API keys are configured in the backend .env file.")
        st.divider()
        st.markdown("**Backend providers**")

        if LLMService:
            try:
                provider_status = LLMService.provider_status()
                for item in provider_status:
                    configured = item.get("configured", False)
                    status_symbol = "✓" if configured else "○"
                    provider_name = item.get("provider", "")
                    provider_model = item.get("model", "")
                    authentication = item.get("authentication", "")

                    st.markdown(
                        f"""
                        <div class="provider-card">
                        <strong>{status_symbol} {PROVIDER_LABELS.get(provider_name, provider_name)}</strong><br>
                        <small>Model: {provider_model}<br>Auth: {authentication}</small>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
            except Exception as exc:
                st.warning(f"Could not read provider status: {exc}")
        else:
            st.caption("LLMService status is unavailable in the dashboard process.")

    st.divider()
    st.caption(f"Backend: {st.session_state.get('api_base')}")
    st.caption(f"Provider: {PROVIDER_LABELS.get(provider, provider)}")
    st.caption(f"Model: {model_name}")

    if st.session_state.get("last_analysis"):
        st.divider()
        if st.button("Start Over", use_container_width=True):
            reset_all()
            st.rerun()


# ==============================================================================
# ACTIVE API CONFIG
# ==============================================================================

ACTIVE_API_BASE = st.session_state.get("api_base", DEFAULT_API_BASE).strip().rstrip("/")

ENDPOINTS = endpoints(ACTIVE_API_BASE)


# ==============================================================================
# HEADER & STEPPER
# ==============================================================================

st.title("🎯 AI Resume & CV Matcher Engine")
st.caption(
    "Analyze ATS compatibility, identify skill gaps, and generate job-tailored resumes."
)

has_analysis = bool(st.session_state.get("last_analysis"))
step = 2 if has_analysis else 1
step1_class = "done" if step > 1 else "active"
step2_class = "active" if step == 2 else ""

st.markdown(
    f"""
    <span class="step-badge {step1_class}">1</span>
    Upload & Analyze
    &nbsp;&nbsp;→&nbsp;&nbsp;
    <span class="step-badge {step2_class}">2</span>
    Review & Export
    """,
    unsafe_allow_html=True,
)

st.write("")


# ==============================================================================
# STEP 1 — INPUT
# ==============================================================================

with st.container(border=True):
    col1, col2 = st.columns(2)

    with col1:
        job_desc = st.text_area(
            "Target Job Description",
            height=240,
            placeholder="Paste the complete job description here...",
            value=st.session_state.get("job_desc", ""),
        )

    with col2:
        uploaded_file = st.file_uploader(
            "Upload Current Resume",
            type=["pdf", "docx", "txt"],
        )

        if uploaded_file:
            size_kb = len(uploaded_file.getvalue()) / 1024
            st.caption(f"📎 {uploaded_file.name} · {size_kb:.0f} KB")

    ready = bool(uploaded_file and job_desc.strip())

    if not ready:
        st.caption("Add both a resume and job description to continue.")

    if st.button(
        "🚀 Analyze ATS Compatibility & Skill Gaps",
        use_container_width=True,
        type="primary",
        disabled=not ready,
    ):
        file_bytes = uploaded_file.getvalue()

        st.session_state["uploaded_file_data"] = (
            uploaded_file.name,
            file_bytes,
            uploaded_file.type or "application/octet-stream",
        )
        st.session_state["job_desc"] = job_desc

        data = {
            "job_description": job_desc,
            "provider": provider,
        }

        files = {
            "resume_file": (
                uploaded_file.name,
                file_bytes,
                uploaded_file.type or "application/octet-stream",
            )
        }

        response = api_request(
            ENDPOINTS["analyze"],
            data=data,
            files=files,
            spinner_text=f"Running ATS analysis with {provider.upper()}...",
        )

        if response is None:
            st.error(f"Could not reach FastAPI backend at {ACTIVE_API_BASE}")
        elif response.status_code == 200:
            try:
                result = response.json()
                st.session_state["last_analysis"] = result
                st.rerun()
            except Exception as exc:
                st.error(f"Backend returned invalid JSON: {exc}")
        else:
            st.error(f"Analysis failed: {error_detail(response)}")


# ==============================================================================
# BACKEND PROCESSING LOG
# ==============================================================================

with st.expander("🛠️ Backend processing log", expanded=False):
    col_log1, col_log2 = st.columns([8, 2])

    with col_log2:
        if st.button("Clear Logs", key="clear_logs_btn", use_container_width=True):
            if LLMService:
                LLMService.clear_logs()
                st.rerun()

    st.caption("Recent LLM requests recorded by the backend.")

    logs = LLMService.recent_logs(50) if LLMService else []

    if logs:
        rows = [
            {
                "Time (UTC)": event.get("timestamp", ""),
                "Event": event.get("event", ""),
                "Provider": event.get("provider", ""),
                "Model": event.get("model", ""),
                "Status": event.get("status", ""),
                "Duration ms": event.get("duration_ms", ""),
                "Prompt chars": event.get("prompt_chars", ""),
                "Response chars": event.get("response_chars", ""),
                "Error": event.get("error", ""),
            }
            for event in reversed(logs)
        ]

        st.dataframe(rows, use_container_width=True, hide_index=True)
        st.caption(f"Log file: {LOG_PATH}")
    else:
        st.info("No backend LLM processing events have been recorded yet.")


# ==============================================================================
# STEP 2 — RESULTS
# ==============================================================================

if st.session_state.get("last_analysis"):
    result = st.session_state["last_analysis"]
    st.write("")

    with st.container(border=True):
        st.subheader("📊 Analysis Results")

        m1, m2, m3 = st.columns(3)

        with m1:
            score_block("ATS Match Score", result.get("ats_match_score", 0))

        with m2:
            score_block("Keyword Density", result.get("keyword_density_score", 0))

        with m3:
            st.markdown(
                '<div class="metric-label">Missing Skills</div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                f'<div class="metric-big">{len(result.get("missing_skills", []))}</div>',
                unsafe_allow_html=True,
            )

        st.divider()

        sc1, sc2 = st.columns(2)

        with sc1:
            st.markdown("**Matching Skills**")
            chips(result.get("matching_skills", []), "good")

        with sc2:
            st.markdown("**Missing Skills**")
            chips(result.get("missing_skills", []), "bad")

    # ==========================================================================
    # IMPROVEMENTS
    # ==========================================================================

    with st.container(border=True):
        st.subheader("💡 Actionable Improvements")

        suggestions = result.get("improvement_suggestions", [])
        if suggestions:
            for suggestion in suggestions:
                if suggestion.startswith("🤖"):
                    st.markdown(suggestion)
                else:
                    st.info(suggestion)
        else:
            st.info("No improvement suggestions were returned.")

    # ==========================================================================
    # FORMAT RECOMMENDATION & LANGUAGE GUARDRAIL ALERT
    # ==========================================================================

    recommendation = result.get("recommendation") or {}
    recommended_format = recommendation.get("recommended_format", "german_corporate")

    if recommendation:
        with st.container(border=True):
            st.subheader("🧭 CV Format Recommendation")

            label = recommendation.get("label", "Standard")
            reason = recommendation.get("reason", "")
            language_mismatch = recommendation.get("language_mismatch", False)

            if language_mismatch:
                st.warning(f"**{label}** — {reason}")
            else:
                st.success(f"**{label}** — {reason}")

    # ==========================================================================
    # DOWNLOAD REPORT
    # ==========================================================================

    with st.expander("📥 Download Analysis Report", expanded=False):
        export_col1, export_col2 = st.columns(2)

        with export_col1:
            st.download_button(
                "📄 Download TXT",
                data=generate_txt_export(result),
                file_name="resume_analysis_report.txt",
                mime="text/plain",
                use_container_width=True,
            )

        with export_col2:
            st.download_button(
                "📝 Download DOCX",
                data=generate_docx_export(result),
                file_name="resume_analysis_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

    # ==========================================================================
    # GENERATION TABS
    # ==========================================================================

    st.divider()
    st.header("✨ Generate a Complete Tailored CV")

    tab1, tab2 = st.tabs(
        [
            "📄 Standard ATS Resume",
            "🇩🇪 German Lebenslauf",
        ]
    )

    with tab1:
        st.caption("Clean single-column Word document optimized for ATS systems.")

        if st.button(
            "🪄 Build Tailored Resume",
            use_container_width=True,
            type="primary",
        ):
            if not build_files_payload():
                st.error(
                    "Resume file is no longer available. Please upload it again."
                )
            else:
                data = {
                    "job_description": st.session_state["job_desc"],
                    "provider": provider,
                }

                response = api_request(
                    ENDPOINTS["full_docx"],
                    data=data,
                    files=build_files_payload(),
                    spinner_text=f"Building tailored resume with {provider.upper()}...",
                )

                if response and response.status_code == 200:
                    st.success("Resume ready.")
                    st.download_button(
                        "📥 Download Resume (.DOCX)",
                        response.content,
                        "Optimized_Tailored_Resume.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                else:
                    st.error(f"Generation failed: {error_detail(response)}")

    with tab2:
        st.caption("Generate a German-style Lebenslauf as PDF or raw LaTeX source.")

        template_options = list(TEMPLATE_LABELS.keys())
        default_index = (
            template_options.index(recommended_format)
            if recommended_format in template_options
            else 0
        )

        selected_layout = st.selectbox(
            "Layout Style",
            options=template_options,
            index=default_index,
            format_func=lambda value: TEMPLATE_LABELS[value],
        )

        common_data = {
            "job_description": st.session_state["job_desc"],
            "layout_style": selected_layout,
            "template_style": selected_layout,
            "provider": provider,
        }

        b1, b2 = st.columns(2)

        with b1:
            if st.button(
                "📄 Build PDF",
                use_container_width=True,
                type="primary",
            ):
                if not build_files_payload():
                    st.error(
                        "Resume file is no longer available. Please upload it again."
                    )
                else:
                    response = api_request(
                        ENDPOINTS["german_pdf"],
                        data=common_data,
                        files=build_files_payload(),
                        spinner_text=f"Generating PDF ({selected_layout})...",
                    )

                    if response and response.status_code == 200:
                        st.success("PDF ready.")
                        st.download_button(
                            "📥 Download PDF",
                            response.content,
                            f"Lebenslauf_{selected_layout}.pdf",
                            "application/pdf",
                            use_container_width=True,
                        )
                    else:
                        st.error(f"PDF generation failed: {error_detail(response)}")

        with b2:
            if st.button(
                "🛠️ Build TEX Source",
                use_container_width=True,
            ):
                if not build_files_payload():
                    st.error(
                        "Resume file is no longer available. Please upload it again."
                    )
                else:
                    response = api_request(
                        ENDPOINTS["german_tex"],
                        data=common_data,
                        files=build_files_payload(),
                        spinner_text="Generating LaTeX source...",
                    )

                    if response and response.status_code == 200:
                        st.success("TEX file ready.")
                        st.download_button(
                            "📥 Download TEX",
                            response.content,
                            f"Lebenslauf_{selected_layout}.tex",
                            "text/plain",
                            use_container_width=True,
                        )
                    else:
                        st.error(f"TEX generation failed: {error_detail(response)}")