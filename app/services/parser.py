import io
from pypdf import PdfReader
from docx import Document
from fastapi import UploadFile, HTTPException


async def extract_text_from_file(file: UploadFile) -> str:
    try:
        # Read file bytes
        content = await file.read()
        # Reset file pointer so other downstream functions can re-read if needed
        await file.seek(0)
    except Exception as e:
        raise HTTPException(
            status_code=400, detail=f"Failed to read upload stream: {str(e)}"
        )

    filename = (file.filename or "").lower()

    # 1. PDF Files
    if filename.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted.strip())
            return "\n\n".join(pages_text).strip()
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to parse PDF file: {str(e)}"
            )

    # 2. DOCX Files (Paragraphs + Tables)
    elif filename.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(content))
            extracted_blocks = []

            # Extract standard paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    extracted_blocks.append(text)

            # Extract content from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in extracted_blocks:
                            extracted_blocks.append(cell_text)

            return "\n\n".join(extracted_blocks).strip()
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to parse DOCX file: {str(e)}"
            )

    # 3. Plain Text Files
    elif filename.endswith(".txt"):
        try:
            return content.decode("utf-8", errors="ignore").strip()
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to parse TXT file: {str(e)}"
            )

    # 4. Fallback for Unsupported Types
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload a PDF, DOCX, or TXT file.",
        )
